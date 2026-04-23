"""Local parsers for Office/CSV/HTML files using lightweight libraries."""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

from .base import BaseParser, EmptyContentError

logger = logging.getLogger(__name__)

_ENCODINGS = ["utf-8", "utf-8-sig", "latin-1", "cp1252"]


def _escape_gfm_cell(value: str) -> str:
    """Escape a cell value so embedded pipes don't break the GFM table.

    Strips leading/trailing whitespace (matches historical cell handling) and
    backslash-escapes every literal ``|`` so markdown viewers don't treat it
    as a column separator.
    """
    return value.strip().replace("|", r"\|")


def _format_gfm_table(rows: list[list[str]]) -> str:
    """Render rows as a GFM table string.

    First row is treated as the header; a ``|---|`` delimiter row of matching
    width is inserted after it. Cells are escaped via ``_escape_gfm_cell``.
    Returns a single string joined by ``\\n`` — no trailing newline.

    Empty ``rows`` raises ``ValueError``; callers should skip emission if a
    sheet has no non-empty rows.
    """
    if not rows:
        raise ValueError("_format_gfm_table requires at least one row")
    col_count = max(len(r) for r in rows)

    def _render(cells: list[str]) -> str:
        # Pad short rows so every row has ``col_count`` columns.
        padded = cells + [""] * (col_count - len(cells))
        return "| " + " | ".join(_escape_gfm_cell(c) for c in padded) + " |"

    delimiter = "|" + "|".join(["---"] * col_count) + "|"
    return "\n".join([_render(rows[0]), delimiter, *(_render(r) for r in rows[1:])])


def _decode_with_fallback(raw_bytes: bytes) -> str:
    """Decode bytes trying common encodings, falling back to lossy UTF-8."""
    for enc in _ENCODINGS:
        try:
            return raw_bytes.decode(enc)
        except (UnicodeDecodeError, ValueError):
            continue
    return raw_bytes.decode("utf-8", errors="replace")


class LocalDocxParser(BaseParser):
    """Parser for DOCX files using python-docx (free, local)."""

    name = "local_docx"

    @property
    def is_available(self) -> bool:
        try:
            import docx  # noqa: F401
            return True
        except ImportError:
            return False

    async def parse(self, file_path: Path) -> str:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            from docx import Document

            doc = Document(str(file_path))
            parts: list[str] = []

            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)

            for table in doc.tables:
                rows: list[str] = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    rows.append(" | ".join(cells))
                if rows:
                    parts.append("\n".join(rows))

            content = "\n\n".join(parts)

            if not content or not content.strip():
                raise EmptyContentError(f"python-docx produced no content for {file_path}")

            logger.info(
                f"Local DOCX parser extracted {len(content)} chars from {file_path.name}"
            )
            return content

        except ImportError:
            raise ValueError(
                "python-docx is not installed. Install with: pip install python-docx"
            )
        except EmptyContentError:
            raise
        except Exception as e:
            raise ValueError(f"Local DOCX parsing failed for {file_path}: {e}") from e


class LocalXlsxParser(BaseParser):
    """Parser for XLSX files using openpyxl (free, local)."""

    name = "local_xlsx"

    @property
    def is_available(self) -> bool:
        try:
            import openpyxl  # noqa: F401
            return True
        except ImportError:
            return False

    async def parse(self, file_path: Path) -> str:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        try:
            from openpyxl import load_workbook

            wb = load_workbook(str(file_path), read_only=True, data_only=True)
            parts: list[str] = []

            for sheet_name in wb.sheetnames:
                ws = wb[sheet_name]
                parts.append(f"## {sheet_name}")

                rows: list[list[str]] = []
                for row in ws.iter_rows(values_only=True):
                    cells = [str(cell) if cell is not None else "" for cell in row]
                    if any(c for c in cells):
                        rows.append(cells)

                if rows:
                    parts.append(_format_gfm_table(rows))

            wb.close()

            content = "\n\n".join(parts)

            if not content or not content.strip():
                raise EmptyContentError(f"openpyxl produced no content for {file_path}")

            logger.info(
                f"Local XLSX parser extracted {len(content)} chars from {file_path.name}"
            )
            return content

        except ImportError:
            raise ValueError(
                "openpyxl is not installed. Install with: pip install openpyxl"
            )
        except EmptyContentError:
            raise
        except Exception as e:
            raise ValueError(f"Local XLSX parsing failed for {file_path}: {e}") from e


class LocalCsvParser(BaseParser):
    """Parser for CSV files using stdlib csv module."""

    name = "local_csv"

    @property
    def is_available(self) -> bool:
        return True

    async def parse(self, file_path: Path) -> str:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        text = _decode_with_fallback(file_path.read_bytes())

        # Some exports have binary garbage appended after the CSV body (NULs +
        # high-ASCII). That trailer produces unquoted bare \r inside a field,
        # which csv.reader rejects ("new-line character seen in unquoted
        # field"). Drop the trailer at the first NUL, then normalize line
        # endings so any remaining lone \r doesn't trip the tokenizer.
        nul = text.find("\x00")
        if nul >= 0:
            text = text[:nul]
        text = text.replace("\r\n", "\n").replace("\r", "\n")

        try:
            reader = csv.reader(io.StringIO(text))
            rows: list[list[str]] = []
            for row in reader:
                cells = [cell.strip() for cell in row]
                if any(c for c in cells):
                    rows.append(cells)

            content = _format_gfm_table(rows) if rows else ""
        except EmptyContentError:
            raise
        except Exception as e:
            raise ValueError(f"Local CSV parsing failed for {file_path}: {e}") from e

        if not content or not content.strip():
            raise EmptyContentError(
                f"CSV parser produced no content for {file_path}"
            )

        logger.info(
            f"Local CSV parser extracted {len(content)} chars from {file_path.name}"
        )
        return content


class LocalHtmlParser(BaseParser):
    """Parser for HTML files using html2text or regex fallback."""

    name = "local_html"

    @property
    def is_available(self) -> bool:
        return True

    async def parse(self, file_path: Path) -> str:
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        html = _decode_with_fallback(file_path.read_bytes())

        try:
            import html2text
            converter = html2text.HTML2Text()
            converter.ignore_links = False
            converter.ignore_images = True
            converter.body_width = 0
            content = converter.handle(html)
        except ImportError:
            # Fallback: use EMLParser's html_to_plain_text method
            from .eml_parser import EMLParser
            parser = EMLParser.__new__(EMLParser)
            content = parser.html_to_plain_text(html)

        if not content or not content.strip():
            raise EmptyContentError(
                f"HTML parser produced no content for {file_path}"
            )

        logger.info(
            f"Local HTML parser extracted {len(content)} chars from {file_path.name}"
        )
        return content
