"""Tests for LocalXlsxParser + LocalCsvParser GFM-table emission.

Covers the forward fix for the broken-pipe-table bug: both parsers now emit
valid GitHub-Flavored Markdown tables (leading/trailing ``|``, a delimiter
row after the header, and escaped ``|`` inside cell values) instead of the
legacy ``cell | cell | cell`` form that markdown viewers collapsed into a
soft-wrapped paragraph.
"""

from __future__ import annotations

from unittest.mock import MagicMock, patch

import pytest

from mtss.parsers.local_office_parser import (
    LocalCsvParser,
    LocalDocxParser,
    LocalXlsxParser,
    _escape_gfm_cell,
    _format_gfm_table,
)


class TestFormatGfmTableHelper:
    """Unit-level coverage for the shared GFM-rendering helper so both parsers
    inherit the same guarantees."""

    def test_emits_header_delimiter_and_data_rows(self):
        out = _format_gfm_table([["A", "B"], ["1", "2"], ["3", "4"]])
        assert out.splitlines() == [
            "| A | B |",
            "|---|---|",
            "| 1 | 2 |",
            "| 3 | 4 |",
        ]

    def test_single_row_still_gets_delimiter(self):
        out = _format_gfm_table([["only"]])
        assert out.splitlines() == ["| only |", "|---|"]

    def test_escapes_literal_pipe_in_cell(self):
        out = _format_gfm_table([["head"], ["a|b"]])
        assert r"| a\|b |" in out
        # Escaped pipe doesn't introduce an extra column — delimiter stays 1-wide.
        assert out.splitlines()[1] == "|---|"

    def test_empty_cells_render_as_empty_segments(self):
        out = _format_gfm_table([["a", "b", "c"], ["1", "", "3"]])
        assert out.splitlines() == [
            "| a | b | c |",
            "|---|---|---|",
            "| 1 |  | 3 |",
        ]

    def test_widens_header_to_max_data_row(self):
        """When a data row is wider than the header (spreadsheets with
        merged/empty header cells), the table widens to the max row width
        and short rows pad with empty cells rather than truncating data."""
        out = _format_gfm_table([["h1", "h2"], ["a", "b", "c"]])
        assert out.splitlines() == [
            "| h1 | h2 |  |",
            "|---|---|---|",
            "| a | b | c |",
        ]

    def test_empty_rows_raise(self):
        with pytest.raises(ValueError):
            _format_gfm_table([])


class TestEscapeGfmCell:
    def test_strips_surrounding_whitespace(self):
        assert _escape_gfm_cell("  hello  ") == "hello"

    def test_escapes_pipe(self):
        assert _escape_gfm_cell("a|b") == r"a\|b"

    def test_preserves_interior_whitespace(self):
        assert _escape_gfm_cell("a  b") == "a  b"


class TestLocalXlsxParserGFM:
    """``LocalXlsxParser`` emits valid GFM tables per sheet."""

    @pytest.fixture
    def fake_workbook(self):
        def _build(sheets: dict[str, list[list]]):
            wb = MagicMock()
            wb.sheetnames = list(sheets.keys())

            def _sheet(name):
                ws = MagicMock()
                ws.iter_rows.return_value = [tuple(row) for row in sheets[name]]
                return ws

            wb.__getitem__.side_effect = _sheet
            wb.close = MagicMock()
            return wb

        return _build

    @pytest.mark.asyncio
    async def test_header_plus_data_rows_render_as_gfm(self, tmp_path, fake_workbook):
        xlsx = tmp_path / "sheet.xlsx"
        xlsx.write_bytes(b"x")
        wb = fake_workbook({
            "Sheet1": [
                ["Name", "Value"],
                ["alpha", 1],
                ["beta", 2],
            ]
        })
        with patch("openpyxl.load_workbook", return_value=wb):
            out = await LocalXlsxParser().parse(xlsx)

        lines = out.splitlines()
        assert "## Sheet1" in lines
        idx = lines.index("| Name | Value |")
        assert lines[idx + 1] == "|---|---|"
        assert "| alpha | 1 |" in lines
        assert "| beta | 2 |" in lines

    @pytest.mark.asyncio
    async def test_cell_with_pipe_is_escaped(self, tmp_path, fake_workbook):
        xlsx = tmp_path / "sheet.xlsx"
        xlsx.write_bytes(b"x")
        wb = fake_workbook({
            "S": [
                ["col"],
                ["foo|bar"],
            ]
        })
        with patch("openpyxl.load_workbook", return_value=wb):
            out = await LocalXlsxParser().parse(xlsx)

        assert r"| foo\|bar |" in out
        # The escaped pipe doesn't inflate the column count.
        assert "|---|" in out.splitlines()

    @pytest.mark.asyncio
    async def test_empty_cells_render_as_empty_segments(self, tmp_path, fake_workbook):
        xlsx = tmp_path / "sheet.xlsx"
        xlsx.write_bytes(b"x")
        wb = fake_workbook({
            "S": [
                ["a", "b", "c"],
                ["1", None, "3"],
            ]
        })
        with patch("openpyxl.load_workbook", return_value=wb):
            out = await LocalXlsxParser().parse(xlsx)

        assert "| a | b | c |" in out
        assert "|---|---|---|" in out
        assert "| 1 |  | 3 |" in out

    @pytest.mark.asyncio
    async def test_blank_rows_are_skipped(self, tmp_path, fake_workbook):
        xlsx = tmp_path / "sheet.xlsx"
        xlsx.write_bytes(b"x")
        wb = fake_workbook({
            "S": [
                ["h1", "h2"],
                [None, None],  # entirely blank — must drop
                ["a", "b"],
            ]
        })
        with patch("openpyxl.load_workbook", return_value=wb):
            out = await LocalXlsxParser().parse(xlsx)

        # No empty data row rendered between header and data.
        lines = [l for l in out.splitlines() if l.startswith("|") and not l.startswith("|---")]
        assert lines == ["| h1 | h2 |", "| a | b |"]

    @pytest.mark.asyncio
    async def test_multiple_sheets_each_get_their_own_table(self, tmp_path, fake_workbook):
        xlsx = tmp_path / "sheet.xlsx"
        xlsx.write_bytes(b"x")
        wb = fake_workbook({
            "Alpha": [["x"], ["1"]],
            "Beta": [["y"], ["2"]],
        })
        with patch("openpyxl.load_workbook", return_value=wb):
            out = await LocalXlsxParser().parse(xlsx)

        assert "## Alpha" in out
        assert "## Beta" in out
        assert out.count("|---|") == 2


class TestLocalCsvParserGFM:
    """``LocalCsvParser`` emits valid GFM tables."""

    @pytest.mark.asyncio
    async def test_header_plus_data_rows_render_as_gfm(self, tmp_path):
        p = tmp_path / "data.csv"
        p.write_bytes(b'"name","value"\n"a",1\n"b",2\n')
        out = await LocalCsvParser().parse(p)
        assert out.splitlines() == [
            "| name | value |",
            "|---|---|",
            "| a | 1 |",
            "| b | 2 |",
        ]

    @pytest.mark.asyncio
    async def test_cell_with_pipe_is_escaped(self, tmp_path):
        p = tmp_path / "pipes.csv"
        p.write_bytes(b'"col"\n"foo|bar"\n')
        out = await LocalCsvParser().parse(p)
        assert r"| foo\|bar |" in out
        # Still a single-column table despite the embedded pipe.
        assert out.splitlines()[1] == "|---|"

    @pytest.mark.asyncio
    async def test_empty_cells_render_as_empty_segments(self, tmp_path):
        p = tmp_path / "gaps.csv"
        p.write_bytes(b"a,b,c\n1,,3\n")
        out = await LocalCsvParser().parse(p)
        assert "| a | b | c |" in out
        assert "|---|---|---|" in out
        assert "| 1 |  | 3 |" in out

    @pytest.mark.asyncio
    async def test_single_row_still_valid_gfm(self, tmp_path):
        """Single-row sheets/CSVs still get a delimiter row — the header
        alone is a valid one-row GFM table."""
        p = tmp_path / "onerow.csv"
        p.write_bytes(b"only\n")
        out = await LocalCsvParser().parse(p)
        assert out.splitlines() == ["| only |", "|---|"]


class TestLocalDocxParserGFM:
    """``LocalDocxParser`` emits valid GFM tables for embedded tables while
    preserving surrounding paragraph text."""

    @staticmethod
    def _write_docx(path, *, paragraphs: list[str], tables: list[list[list[str]]]):
        """Build a minimal .docx via python-docx with paragraphs then tables."""
        from docx import Document

        doc = Document()
        for para in paragraphs:
            doc.add_paragraph(para)
        for table_rows in tables:
            if not table_rows:
                continue
            cols = max(len(r) for r in table_rows)
            tbl = doc.add_table(rows=len(table_rows), cols=cols)
            for r_idx, row in enumerate(table_rows):
                for c_idx in range(cols):
                    tbl.rows[r_idx].cells[c_idx].text = row[c_idx] if c_idx < len(row) else ""
        doc.save(str(path))

    @pytest.mark.asyncio
    async def test_header_plus_data_rows_render_as_gfm(self, tmp_path):
        p = tmp_path / "doc.docx"
        self._write_docx(
            p,
            paragraphs=[],
            tables=[[
                ["Name", "Value"],
                ["alpha", "1"],
                ["beta", "2"],
            ]],
        )
        out = await LocalDocxParser().parse(p)
        lines = out.splitlines()
        assert "| Name | Value |" in lines
        idx = lines.index("| Name | Value |")
        assert lines[idx + 1] == "|---|---|"
        assert "| alpha | 1 |" in lines
        assert "| beta | 2 |" in lines

    @pytest.mark.asyncio
    async def test_cell_with_pipe_is_escaped(self, tmp_path):
        p = tmp_path / "pipes.docx"
        self._write_docx(
            p,
            paragraphs=[],
            tables=[[
                ["col"],
                ["foo|bar"],
            ]],
        )
        out = await LocalDocxParser().parse(p)
        assert r"| foo\|bar |" in out
        # Escaped pipe doesn't inflate the column count.
        assert "|---|" in out.splitlines()

    @pytest.mark.asyncio
    async def test_blank_rows_are_skipped(self, tmp_path):
        """Entirely-empty rows in the docx table should be dropped before
        GFM rendering, matching xlsx behaviour — otherwise the delimiter
        row or a data row would render as `|  |  |` noise."""
        p = tmp_path / "blanks.docx"
        self._write_docx(
            p,
            paragraphs=[],
            tables=[[
                ["h1", "h2"],
                ["", ""],  # entirely blank — must drop
                ["a", "b"],
            ]],
        )
        out = await LocalDocxParser().parse(p)
        data_lines = [l for l in out.splitlines() if l.startswith("|") and not l.startswith("|---")]
        assert data_lines == ["| h1 | h2 |", "| a | b |"]

    @pytest.mark.asyncio
    async def test_paragraphs_preserved_above_table(self, tmp_path):
        """Existing docx parser behaviour: paragraph text comes out first,
        then the table. The GFM-fix must not break that ordering or drop
        prose."""
        p = tmp_path / "mixed.docx"
        self._write_docx(
            p,
            paragraphs=["Intro sentence.", "Second paragraph."],
            tables=[[
                ["A", "B"],
                ["1", "2"],
            ]],
        )
        out = await LocalDocxParser().parse(p)
        assert "Intro sentence." in out
        assert "Second paragraph." in out
        # Paragraph text precedes the GFM table.
        assert out.index("Intro sentence.") < out.index("| A | B |")
        # Table is valid GFM.
        lines = out.splitlines()
        idx = lines.index("| A | B |")
        assert lines[idx + 1] == "|---|---|"
        assert "| 1 | 2 |" in lines
